"""
LD-friendly DOCX and HTML exports with clear sections, spacing, and colours.
"""

import html
import io
import json
import zipfile
from typing import Any

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, RGBColor

from concept_map_builder import add_vocabulary_concept_map_to_docx, build_vocabulary_concept_map_svg
from html_exporter import export_tab_html
from structured_renderers import _as_dict, lesson_to_text, vocabulary_to_text, worksheet_to_text

NAVY = RGBColor(0x0B, 0x2E, 0x59)
TEAL = RGBColor(0x00, 0x8C, 0x95)


def _apply_ld_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_after = Pt(10)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=min(level, 3))
    for run in heading.runs:
        run.font.color.rgb = NAVY if level == 1 else TEAL
        run.font.name = "Arial"


def _add_colored_box(doc: Document, label: str, body: str, rgb: RGBColor) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = rgb
    run.font.name = "Arial"
    run.font.size = Pt(14)
    p.add_run(body)


def _add_bullets(doc: Document, items: list) -> None:
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def export_vocabulary_docx(data: Any) -> bytes:
    vocab = _as_dict(data) or {}
    doc = Document()
    _apply_ld_normal_style(doc)
    _add_heading(doc, f"Vocabulary Study — {vocab.get('topic', 'Lesson')}")

    _add_heading(doc, "1. Word Wall", 2)
    for word in vocab.get("word_wall") or []:
        _add_colored_box(
            doc,
            word.get("term", "Term"),
            word.get("definition", ""),
            TEAL,
        )
        visual = word.get("visual_description") or word.get("visual", "")
        if visual:
            doc.add_paragraph(f"Picture: {visual}")

    _add_heading(doc, "2. Flashcards", 2)
    for card in vocab.get("flashcards") or []:
        doc.add_paragraph(
            f"FRONT: {card.get('front', '')}  →  BACK: {card.get('back', '')}"
        )

    _add_heading(doc, "3. Picture Words", 2)
    for row in vocab.get("picture_words") or []:
        doc.add_paragraph(
            f"{row.get('term', '')} — {row.get('color_cue', '')} — "
            f"Draw: {row.get('draw_this', '')}"
        )

    _add_heading(doc, "4. Say · Spell · Use", 2)
    for item in vocab.get("practice") or []:
        doc.add_paragraph(
            f"{item.get('term', '')} ({item.get('pronunciation', '')}) — "
            f"{item.get('sentence_blank', '')}"
        )

    _add_heading(doc, "5. Self-Test", 2)
    self_test = vocab.get("self_test") or {}
    if self_test.get("matching_prompt"):
        doc.add_paragraph(str(self_test["matching_prompt"]))
    for sentence in self_test.get("fill_blanks") or []:
        doc.add_paragraph(str(sentence))

    _add_heading(doc, "6. Quick Reference", 2)
    for row in vocab.get("reference_chart") or []:
        doc.add_paragraph(
            f"{row.get('term', '')}: {row.get('definition', '')} "
            f"(Exam tip: {row.get('exam_tip', '')})"
        )

    add_vocabulary_concept_map_to_docx(doc, vocab)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_worksheet_docx(data: Any) -> bytes:
    sheet = _as_dict(data) or {}
    doc = Document()
    _apply_ld_normal_style(doc)
    header = sheet.get("header") or {}
    _add_heading(doc, "Exam Practice Worksheet")
    doc.add_paragraph(
        f"Subject: {header.get('subject', '')}  |  Topic: {header.get('topic', '')}  |  "
        f"Time: {header.get('time_allowed', '')}  |  Marks: {header.get('total_marks', '')}"
    )
    doc.add_paragraph("Name: _________________________    Date: _________________________")

    _add_heading(doc, "Part A — Short Answer", 2)
    for index, item in enumerate(sheet.get("short_answer") or [], 1):
        doc.add_paragraph(f"Q{index}. ({item.get('marks', 2)} marks) {item.get('question', '')}")
        for _ in range(int(item.get("lines", 3))):
            doc.add_paragraph("_" * 70)

    _add_heading(doc, "Part B — Long Answer", 2)
    for index, item in enumerate(sheet.get("long_answer") or [], 1):
        doc.add_paragraph(f"Q{index}. ({item.get('marks', 6)} marks) {item.get('question', '')}")
        for _ in range(int(item.get("lines", 8))):
            doc.add_paragraph("_" * 70)

    diagram = sheet.get("diagram_question") or {}
    if diagram:
        _add_heading(doc, "Part C — Diagram", 2)
        doc.add_paragraph(f"({diagram.get('marks', 4)} marks) {diagram.get('question', '')}")
        doc.add_paragraph("[Sketch and label the diagram on your answer sheet]")

    _add_heading(doc, "Part D — Vocabulary", 2)
    for index, item in enumerate(sheet.get("vocab_questions") or [], 1):
        doc.add_paragraph(f"{index}. ({item.get('marks', 1)} mark) {item.get('question', '')}")
        doc.add_paragraph("Answer: _________________________________")

    _add_heading(doc, "Part E — Exam Ready Checklist", 2)
    _add_bullets(doc, sheet.get("student_checklist") or [])

    _add_heading(doc, "Teacher Answer Key", 2)
    for item in sheet.get("answer_key") or []:
        doc.add_paragraph(f"{item.get('question_ref', '')}: {item.get('model_answer', '')}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_lesson_docx(data: Any, title: str) -> bytes:
    lesson = _as_dict(data)
    doc = Document()
    _apply_ld_normal_style(doc)
    _add_heading(doc, title)

    if lesson:
        if lesson.get("big_idea"):
            _add_colored_box(doc, "Big Idea", lesson["big_idea"], TEAL)
        for section in lesson.get("sections") or []:
            _add_heading(doc, section.get("title", "Section"), 2)
            body = section.get("body", "")
            for para in body.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        if lesson.get("mermaid_diagram"):
            doc.add_paragraph("[Concept diagram — see online tab or draw from Visual Summary below]")
        for item in lesson.get("visual_summary") or []:
            doc.add_paragraph(f"{item.get('icon', '')} {item.get('color_name', '')}: {item.get('idea', '')}")
    else:
        for para in str(data).split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_tab_docx(title: str, content: Any, spec_id: str) -> bytes:
    if spec_id == "vocabulary":
        return export_vocabulary_docx(content)
    if spec_id == "worksheet":
        return export_worksheet_docx(content)
    if spec_id == "original":
        doc = Document()
        _apply_ld_normal_style(doc)
        _add_heading(doc, title)
        for para in str(content).split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    return export_lesson_docx(content, title)


def export_tab_html(title: str, content: Any, spec_id: str) -> str:
    """Delegate to html_exporter for rich LD-friendly HTML."""
    from html_exporter import export_tab_html as _rich_html

    return _rich_html(title, content, spec_id)


def build_zip_bundle(
    specs: list,
    adaptations: dict,
    lesson_text: str,
    base_name: str,
) -> bytes:
    """ZIP with separate DOCX file per tab — easy to print and share."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for spec in specs:
            if spec["id"] == "original":
                content = lesson_text
            else:
                content = adaptations.get(spec["id"], "")
            safe_title = spec["id"]
            docx_bytes = export_tab_docx(spec["title"], content, spec["id"])
            zf.writestr(f"{base_name}_{safe_title}.docx", docx_bytes)
            html_str = export_tab_html(spec["title"], content, spec["id"])
            zf.writestr(f"{base_name}_{safe_title}.html", html_str)
            if spec["id"] == "vocabulary":
                svg = build_vocabulary_concept_map_svg(content)
                zf.writestr(f"{base_name}_vocabulary_concept_map.svg", svg)
    return buffer.getvalue()
