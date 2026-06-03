"""
Create sample_lesson.docx for testing EduAdapt AI without a real upload.
Run once: python create_sample_lesson.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt


def build_sample_document() -> Document:
    """
    Build a Grade 6 science lesson with clear objectives for analytics testing.

    Returns:
        A python-docx Document ready to save.
    """
    doc = Document()

    title = doc.add_heading("The Water Cycle: How Earth's Water Moves", level=0)
    title.runs[0].font.size = Pt(18)

    doc.add_paragraph("Grade Level: 6 | Subject: Earth Science | Time: 45 minutes")

    doc.add_heading("Learning Objectives", level=1)
    objectives = [
        "Students will explain how water changes state during the water cycle.",
        "Students will identify evaporation, condensation, precipitation, and collection.",
        "Students will describe how the sun drives the water cycle.",
        "Students will create a labeled diagram of the water cycle.",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style="List Bullet")

    doc.add_heading("Essential Question", level=1)
    doc.add_paragraph(
        "How does water travel through Earth's systems, and why does it matter for life?"
    )

    doc.add_heading("Introduction (10 minutes)", level=1)
    doc.add_paragraph(
        "Begin by asking students where they saw water today—faucets, clouds, puddles, "
        "plants. Record responses on the board. Show a short image of a cloud forming "
        "over a lake. Tell students that the same water molecules have been cycling on "
        "Earth for billions of years."
    )

    doc.add_heading("Direct Instruction (15 minutes)", level=1)
    doc.add_paragraph(
        "The water cycle is the continuous movement of water among the atmosphere, "
        "land, and oceans. Energy from the sun heats surface water. Evaporation occurs "
        "when liquid water becomes water vapor. Transpiration adds water vapor from plants. "
        "As vapor rises, it cools and condenses into tiny droplets that form clouds. "
        "When droplets combine and grow heavy, precipitation falls as rain, snow, sleet, "
        "or hail. Collection happens when water gathers in rivers, lakes, groundwater, "
        "and oceans—and the cycle begins again."
    )

    doc.add_heading("Guided Practice (10 minutes)", level=1)
    doc.add_paragraph(
        "In pairs, students receive vocabulary cards: evaporation, condensation, "
        "precipitation, collection, transpiration. They sequence the steps on a desk "
        "timeline and justify their order to another pair."
    )

    doc.add_heading("Independent Practice (10 minutes)", level=1)
    doc.add_paragraph(
        "Each student draws and labels a water cycle diagram. They must include the sun, "
        "a body of water, a cloud, arrows showing movement, and at least one example of "
        "human impact (e.g., irrigation, pollution runoff)."
    )

    doc.add_heading("Assessment", level=1)
    doc.add_paragraph(
        "Exit ticket: Name two processes in the water cycle and explain what causes each. "
        "Success criteria: correct vocabulary, clear cause-and-effect, complete sentences."
    )

    doc.add_heading("Differentiation Preview", level=1)
    doc.add_paragraph(
        "Support: word bank and sentence frames. Extend: research how climate change "
        "affects regional precipitation patterns and write a short paragraph."
    )

    return doc


def main() -> None:
    """Write sample_lesson.docx to the samples folder."""
    samples_dir = Path(__file__).parent / "samples"
    samples_dir.mkdir(exist_ok=True)
    output_path = samples_dir / "sample_lesson.docx"

    build_sample_document().save(output_path)
    print(f"Created: {output_path}")


if __name__ == "__main__":
    main()
