"""
Build a printable vocabulary concept map (SVG + Word table) in EduAdapt teal/navy branding.
"""

from __future__ import annotations

import html
import json
import math
from typing import Any


def _coerce_dict(content: Any) -> dict | None:
    if isinstance(content, dict):
        return content
    if isinstance(content, str):
        text = content.strip()
        if text.startswith("{"):
            try:
                parsed = json.loads(text)
                if isinstance(parsed, dict):
                    return parsed
            except json.JSONDecodeError:
                pass
    return None

NAVY = "#0B2E59"
TEAL = "#008C95"
LIGHT_TEAL = "#e6f7f8"
LIGHT_BLUE = "#e3f2fd"
FONT = "Lexend, Arial, Verdana, sans-serif"


def _topic_and_terms(vocab: dict) -> tuple[str, list[str]]:
    try:
        from engines.lesson_composition_engine.vocab_quality import clean_topic, is_junk_term
    except Exception:
        clean_topic = lambda t, fallback="Lesson Topic": t or fallback  # noqa: E731
        is_junk_term = lambda t: False  # noqa: E731

    topic = clean_topic((vocab.get("topic") or "Lesson Topic").strip())
    terms: list[str] = []
    for word in vocab.get("word_wall") or []:
        term = (word.get("term") or "").strip()
        if term and term not in terms and not is_junk_term(term):
            terms.append(term)
    if len(terms) < 3:
        for row in vocab.get("reference_chart") or []:
            term = (row.get("term") or "").strip()
            if term and term not in terms and not is_junk_term(term):
                terms.append(term)
    if len(terms) < 1:
        for card in vocab.get("flashcards") or []:
            term = (card.get("front") or card.get("term") or "").strip()
            if term and term not in terms and not is_junk_term(term):
                terms.append(term)
    if len(terms) < 1:
        for row in vocab.get("picture_words") or []:
            term = (row.get("term") or "").strip()
            if term and term not in terms and not is_junk_term(term):
                terms.append(term)
    return topic, terms[:8]


def _wrap_label(text: str, max_len: int = 14) -> list[str]:
    words = text.split()
    if not words:
        return [text[:max_len]]
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        if len(current) + 1 + len(word) <= max_len:
            current += f" {word}"
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines[:2]


def _hub_size(topic: str) -> tuple[int, int]:
    lines = _wrap_label(topic, 14)
    hub_w = min(280, max(150, max(len(line) for line in lines) * 9 + 36))
    hub_h = max(48, 22 + len(lines) * 22)
    return hub_w, hub_h


def _box_size(term: str) -> tuple[int, int, list[str]]:
    lines = _wrap_label(term, 13)
    box_w = min(170, max(96, max(len(line) for line in lines) * 8 + 28))
    box_h = max(44, 18 + len(lines) * 18)
    return box_w, box_h, lines


def _hub_border_point(cx: float, cy: float, hw: float, hh: float, angle: float) -> tuple[float, float]:
    """Intersection of ray from hub centre with hub rectangle border."""
    dx = math.cos(angle)
    dy = math.sin(angle)
    if abs(dx) < 1e-9 and abs(dy) < 1e-9:
        return cx, cy
    scale = float("inf")
    if abs(dx) > 1e-9:
        scale = min(scale, hw / abs(dx))
    if abs(dy) > 1e-9:
        scale = min(scale, hh / abs(dy))
    return cx + dx * scale, cy + dy * scale


def build_relationship_concept_map_svg(vocab: Any, *, width: int = 920, height: int = 520) -> str:
    """Publication-quality vertical relationship map from CLG edges (not a keyword cloud)."""
    data = _coerce_dict(vocab) or {}
    topic = str(data.get("topic") or "Lesson").strip()
    edges = list(data.get("concept_map_edges") or data.get("relationships") or [])
    terms: list[str] = []
    if edges:
        # Prefer ordered path from relationship labels / concept chain
        for edge in edges:
            if not isinstance(edge, dict):
                continue
            label = str(edge.get("label") or "")
            for part in label.replace(" leads to ", "→").split("→"):
                name = part.strip()
                if name and name not in terms:
                    terms.append(name)
        if len(terms) < 2:
            topic2, wall_terms = _topic_and_terms(data)
            topic = topic or topic2
            terms = wall_terms
    else:
        topic, terms = _topic_and_terms(data)

    if len(terms) < 2:
        return _build_hub_spoke_concept_map_svg(data)

    terms = terms[:8]
    colours = [TEAL, NAVY, "#0ea5e9", "#059669", "#0284c7", "#334155", "#0891b2", "#0f766e"]
    cx = width / 2
    top = 88
    box_w, box_h = 280, 56
    gap = 18
    needed_h = top + len(terms) * (box_h + gap) + 40
    height = max(height, needed_h)

    parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" '
        f'viewBox="0 0 {width} {height}" role="img" aria-label="Concept relationship map">'
        f"<defs>"
        f'<marker id="cmap-arrow" markerWidth="10" markerHeight="10" refX="8" refY="3" orient="auto">'
        f'<path d="M0,0 L8,3 L0,6 Z" fill="{NAVY}"/></marker>'
        f'<filter id="cmap-shadow" x="-20%" y="-20%" width="140%" height="140%">'
        f'<feDropShadow dx="0" dy="4" stdDeviation="5" flood-color="{NAVY}" flood-opacity=".12"/>'
        f"</filter></defs>"
        f'<rect width="{width}" height="{height}" rx="22" fill="#f8fbfd"/>'
        f'<rect x="24" y="18" width="{width-48}" height="50" rx="14" fill="{NAVY}"/>'
        f'<text x="{cx}" y="50" text-anchor="middle" font-family="{FONT}" font-size="20" '
        f'font-weight="700" fill="#fff">{html.escape(topic[:64])}</text>'
        f'<text x="{cx}" y="82" text-anchor="middle" font-family="{FONT}" font-size="12" '
        f'fill="#475569">Concept relationships (not a keyword cloud)</text>',
    ]
    y = top
    for i, term in enumerate(terms):
        fill = colours[i % len(colours)]
        x = cx - box_w / 2
        parts.append(
            f'<rect x="{x}" y="{y}" width="{box_w}" height="{box_h}" rx="16" fill="{fill}" '
            f'filter="url(#cmap-shadow)"/>'
            f'<text x="{cx}" y="{y + box_h/2 + 5}" text-anchor="middle" font-family="{FONT}" '
            f'font-size="16" font-weight="700" fill="#fff">{html.escape(term[:36])}</text>'
        )
        if i < len(terms) - 1:
            parts.append(
                f'<line x1="{cx}" y1="{y + box_h}" x2="{cx}" y2="{y + box_h + gap}" '
                f'stroke="{NAVY}" stroke-width="3" marker-end="url(#cmap-arrow)"/>'
            )
        y += box_h + gap

    # Legend
    parts.append(
        f'<text x="36" y="{height - 18}" font-family="{FONT}" font-size="12" fill="#64748b">'
        f"Hierarchy shows teaching sequence and relationships</text>"
    )
    parts.append("</svg>")
    return "\n".join(parts)


def build_vocabulary_concept_map_svg(vocab: Any) -> str:
    """Prefer CLG relationship maps; fall back to hub-and-spoke for legacy walls."""
    data = _coerce_dict(vocab) or {}
    if data.get("concept_map_edges") or data.get("relationships"):
        return build_relationship_concept_map_svg(data)
    return _build_hub_spoke_concept_map_svg(data)


def _build_hub_spoke_concept_map_svg(vocab: Any) -> str:
    """
    Hub-and-spoke map: topic inside the centre navy box; term boxes arranged
    in a ring around it with labels inside each box (no overlapping text).
    """
    data = _coerce_dict(vocab) or {}
    topic, terms = _topic_and_terms(data)
    if not terms:
        terms = ["Key term 1", "Key term 2", "Key term 3"]

    width, height = 920, 760
    cx, cy = width // 2, height // 2
    topic_lines = _wrap_label(topic, 14)
    hub_w, hub_h = _hub_size(topic)
    hub_x = cx - hub_w // 2
    hub_y = cy - hub_h // 2
    hw, hh = hub_w / 2, hub_h / 2

    count = len(terms)
    radius = max(250, min(310, 210 + count * 10))

    nodes: list[dict] = []
    for index, term in enumerate(terms):
        angle = -math.pi / 2 + (2 * math.pi * index / count)
        tx = cx + radius * math.cos(angle)
        ty = cy + radius * math.sin(angle)
        box_w, box_h, label_lines = _box_size(term)
        bx = tx - box_w / 2
        by = ty - box_h / 2
        ex, ey = _hub_border_point(cx, cy, hw, hh, angle)
        nodes.append(
            {
                "term": term,
                "angle": angle,
                "tx": tx,
                "ty": ty,
                "bx": bx,
                "by": by,
                "box_w": box_w,
                "box_h": box_h,
                "label_lines": label_lines,
                "ex": ex,
                "ey": ey,
                "fill": LIGHT_TEAL if index % 2 == 0 else LIGHT_BLUE,
            }
        )

    parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" '
        f'viewBox="0 0 {width} {height}" role="img" '
        f'aria-label="Concept map for {html.escape(topic)}">',
        f'<defs><marker id="arrow" markerWidth="8" markerHeight="8" refX="6" refY="3" '
        f'orient="auto"><path d="M0,0 L6,3 L0,6 Z" fill="{TEAL}"/></marker></defs>',
        f'<rect width="100%" height="100%" fill="#fafcfd"/>',
    ]

    # Lines first (under boxes).
    for node in nodes:
        target_y = node["by"] + node["box_h"] / 2
        target_x = node["tx"]
        if node["ty"] < cy:
            target_y = node["by"] + node["box_h"]
        elif node["ty"] > cy:
            target_y = node["by"]
        if node["tx"] < cx - 20:
            target_x = node["bx"] + node["box_w"]
        elif node["tx"] > cx + 20:
            target_x = node["bx"]
        parts.append(
            f'<line x1="{node["ex"]:.1f}" y1="{node["ey"]:.1f}" '
            f'x2="{target_x:.1f}" y2="{target_y:.1f}" stroke="{TEAL}" '
            f'stroke-width="2" marker-end="url(#arrow)"/>'
        )

    # Centre hub with topic label inside.
    parts.append(
        f'<rect x="{hub_x}" y="{hub_y}" width="{hub_w}" height="{hub_h}" rx="12" '
        f'fill="{NAVY}" stroke="{TEAL}" stroke-width="2"/>'
    )
    text_start_y = hub_y + 20 + (hub_h - len(topic_lines) * 22) // 2
    for index, line in enumerate(topic_lines):
        ty = text_start_y + index * 22
        parts.append(
            f'<text x="{cx}" y="{ty}" text-anchor="middle" font-family="{FONT}" '
            f'font-size="14" font-weight="700" fill="#ffffff">{html.escape(line)}</text>'
        )

    # Term boxes on top with labels centred inside each box.
    for node in nodes:
        parts.append(
            f'<rect x="{node["bx"]:.1f}" y="{node["by"]:.1f}" '
            f'width="{node["box_w"]:.1f}" height="{node["box_h"]:.1f}" rx="10" '
            f'fill="{node["fill"]}" stroke="{TEAL}" stroke-width="2"/>'
        )
        line_count = len(node["label_lines"])
        label_start = node["ty"] - (line_count - 1) * 9
        for line_index, line in enumerate(node["label_lines"]):
            ly = label_start + line_index * 18
            parts.append(
                f'<text x="{node["tx"]:.1f}" y="{ly:.1f}" text-anchor="middle" '
                f'dominant-baseline="middle" font-family="{FONT}" font-size="12" '
                f'font-weight="600" fill="{NAVY}">{html.escape(line)}</text>'
            )

    parts.append("</svg>")
    return "\n".join(parts)


def _shade_cell(cell, fill_hex: str) -> None:
    """Apply background colour to a Word table cell."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def render_concept_map_streamlit(vocab: Any) -> None:
    """Always-visible hub-and-spoke map in the app (SVG only — section title is elsewhere)."""
    import streamlit as st

    svg = build_vocabulary_concept_map_svg(vocab)
    st.markdown(
        f'<div style="display:flex;justify-content:center;align-items:center;'
        f'max-width:100%;overflow-x:auto;padding:0.5rem 0 1rem 0;">{svg.strip()}</div>',
        unsafe_allow_html=True,
    )


def add_vocabulary_concept_map_to_docx(doc, vocab: Any) -> None:
    """Section 7 — centred flowchart-style concept map for Word print."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    data = _coerce_dict(vocab) or {}
    topic, terms = _topic_and_terms(data)
    if not terms:
        terms = ["Key term 1", "Key term 2", "Key term 3"]

    heading = doc.add_heading("7. Concept Map", level=2)
    for run in heading.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0x0B, 0x2E, 0x59)

    intro = doc.add_paragraph(
        "Study how all vocabulary terms connect to the main topic."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hub = doc.add_table(rows=1, cols=1)
    hub.style = "Table Grid"
    hub_cell = hub.rows[0].cells[0]
    _shade_cell(hub_cell, "0B2E59")
    hub_para = hub_cell.paragraphs[0]
    hub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hub_run = hub_para.add_run(topic)
    hub_run.bold = True
    hub_run.font.name = "Arial"
    hub_run.font.size = Pt(16)
    hub_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph("")

    columns = 3
    rows_needed = (len(terms) + columns - 1) // columns
    grid = doc.add_table(rows=rows_needed, cols=columns)
    grid.style = "Table Grid"

    for index, term in enumerate(terms):
        row = grid.rows[index // columns]
        cell = row.cells[index % columns]
        fill = "E6F7F8" if index % 2 == 0 else "E3F2FD"
        _shade_cell(cell, fill)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(term)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x0B, 0x2E, 0x59)

    note = doc.add_paragraph(
        "Tip: Download the HTML vocabulary file for the full arrow flowchart (best for colour printing)."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in note.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x00, 0x8C, 0x95)


def vocabulary_concept_map_html_block(vocab: Any) -> str:
    """Centered concept map section for HTML export."""
    data = _coerce_dict(vocab) or {}
    svg = build_vocabulary_concept_map_svg(data)
    return (
        '<h2>7. Concept Map</h2>'
        '<p style="text-align:center;color:#0B2E59;">Study how all vocabulary terms connect to the main topic.</p>'
        f'<div class="concept-map">{svg}</div>'
    )
